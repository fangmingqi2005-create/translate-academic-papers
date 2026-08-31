#!/usr/bin/env python3
"""Resumable PDF -> Chinese/bilingual DOCX pipeline.

The translation cache is JSONL so interrupted runs resume without retranslating
completed source IDs.  A local Argos en->zh model is used only as a deterministic
fallback; an AI agent should revise the cache for terminology and academic style
before final delivery.
"""
from __future__ import annotations

import argparse
import hashlib
import importlib
import json
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from pypdf import PdfReader


def setup_argos(runtime: Path, models: Path):
    sys.path.insert(0, str(runtime))
    os.environ["ARGOS_PACKAGES_DIR"] = str(models)
    os.environ["PYTHONIOENCODING"] = "utf-8"
    return importlib.import_module("argostranslate.translate")


def normalize(text: str) -> str:
    text = text.replace("\u00ad", "").replace("\ufeff", "").replace("\r", "\n")
    text = re.sub(r"(?<=\w)-\n(?=[a-z])", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def extract_records(pdf: Path) -> list[dict]:
    try:
        import pymupdf
    except ImportError:
        return extract_records_pypdf(pdf)

    doc = pymupdf.open(pdf)
    records: list[dict] = []
    in_references = False
    for page_no, page in enumerate(doc, 1):
        width, height = page.rect.width, page.rect.height
        raw = []
        for b in page.get_text("blocks", sort=False):
            x0, y0, x1, y1, text = b[:5]
            text = normalize(text.replace("\n", " "))
            if not text or x0 < 20 or y1 > height - 28:
                continue
            if re.search(r"Downloaded from https?://", text, re.I):
                continue
            if re.search(r"PNAS\s+\d{4}.*\d+ of \d+", text, re.I):
                continue
            raw.append((x0, y0, x1, y1, text))

        # Journal pages are predominantly two-column. Preserve top spanning
        # material, then read the left column before the right column.
        top = [b for b in raw if b[1] < 145 and (b[2] - b[0]) > width * .55]
        rest = [b for b in raw if b not in top]
        left = [b for b in rest if (b[0] + b[2]) / 2 < width / 2]
        right = [b for b in rest if (b[0] + b[2]) / 2 >= width / 2]
        ordered = sorted(top, key=lambda b: (b[1], b[0])) + sorted(left, key=lambda b: (b[1], b[0])) + sorted(right, key=lambda b: (b[1], b[0]))
        block_no = 0
        for _, _, _, _, block in ordered:
            if re.fullmatch(r"(?:OPEN ACCESS|RESEARCH ARTICLE.*)", block, re.I):
                continue
            if (
                re.fullmatch(r"References?", block, re.I)
                or block.lower().startswith("references ")
                or (page_no >= 3 and re.match(r"^\s*1\.\s+[A-Z]", block))
            ):
                in_references = True
            block_no += 1
            records.append({"id": f"p{page_no:02d}_b{block_no:03d}", "page": page_no, "kind": "reference" if in_references else "body", "en": block})
    return records


def extract_records_pypdf(pdf: Path) -> list[dict]:
    records: list[dict] = []
    in_references = False
    for page_no, page in enumerate(PdfReader(str(pdf)).pages, 1):
        text = normalize(page.extract_text() or "")
        blocks = [normalize(x) for x in re.split(r"\n\s*\n", text) if normalize(x)]
        if len(blocks) <= 2:
            # Many journal PDFs expose one line per visual line. Rejoin until a
            # sentence ending or a short heading is reached.
            blocks, buf = [], []
            for line in [normalize(x) for x in text.split("\n") if normalize(x)]:
                buf.append(line)
                joined = " ".join(buf)
                if re.search(r"[.!?][\]\)\"']?$", line) or len(joined) > 900 or len(line) < 55:
                    blocks.append(joined)
                    buf = []
            if buf:
                blocks.append(" ".join(buf))
        for block in blocks:
            if (
                re.fullmatch(r"References?", block, re.I)
                or block.lower().startswith("references ")
                or (page_no >= 3 and re.match(r"^\s*1\.\s+[A-Z]", block))
            ):
                in_references = True
            records.append({
                "id": f"p{page_no:02d}_b{len([r for r in records if r['page']==page_no])+1:03d}",
                "page": page_no,
                "kind": "reference" if in_references else "body",
                "en": block,
            })
    return records


def load_cache(path: Path) -> dict[str, dict]:
    out = {}
    if path.exists():
        for line in path.read_text(encoding="utf-8").splitlines():
            if line.strip():
                row = json.loads(line)
                # Last record wins, which allows an agent to append a revised
                # AI-verified translation without destructively rewriting the
                # resumable cache.
                out[row.get("key", row["id"])] = row
    return out


def append_cache(path: Path, record_id: str, key: str, en: str, zh: str, status: str):
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as f:
        f.write(json.dumps({
            "id": record_id,
            "key": key,
            "en": en,
            "zh": zh,
            "status": status,
        }, ensure_ascii=False) + "\n")


def record_key(record: dict) -> str:
    digest = hashlib.sha1(record["en"].encode("utf-8")).hexdigest()[:12]
    return f"{record['id']}:{digest}"


def translate_records(records: list[dict], cache_path: Path, translator) -> dict[str, dict]:
    cache = load_cache(cache_path)
    total = sum(r["kind"] == "body" for r in records)
    done = sum(record_key(r) in cache for r in records if r["kind"] == "body")
    by_id = {}
    for r in records:
        if r["kind"] == "reference":
            continue
        key = record_key(r)
        if key in cache:
            by_id[r["id"]] = cache[key]
            continue
        zh = translator.translate(r["en"], "en", "zh").strip()
        if not zh:
            raise RuntimeError(f"empty translation for {r['id']}")
        row = {"id": r["id"], "key": key, "en": r["en"], "zh": zh, "status": "machine_draft"}
        cache[key] = row
        by_id[r["id"]] = row
        append_cache(cache_path, r["id"], key, r["en"], zh, "machine_draft")
        done += 1
        print(f"translated {done}/{total}: {r['id']}", flush=True)
    return by_id


def set_font(run, name: str, size: float, color=None, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_bookmark(paragraph, name: str, bookmark_id: int):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def base_doc(title: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin = sec.right_margin = Cm(2.54)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(title), "黑体", 18, bold=True)
    return doc


def build_docs(records: list[dict], cache: dict[str, dict], out_cn: Path, out_bi: Path, title_zh: str):
    cn, bi = base_doc(title_zh), base_doc(title_zh + "（逐段中英对照）")
    bookmark_id = 1
    for r in records:
        if r["kind"] == "reference":
            p1 = cn.add_paragraph(r["en"])
            p2 = bi.add_paragraph(r["en"])
            m = re.match(r"\s*(\d+)\.?\s+", r["en"])
            if m:
                add_bookmark(p1, f"ref_{m.group(1)}", bookmark_id); bookmark_id += 1
                add_bookmark(p2, f"ref_{m.group(1)}", bookmark_id); bookmark_id += 1
            continue
        zh = cache[r["id"]]["zh"]
        p = cn.add_paragraph(zh)
        p.paragraph_format.first_line_indent = Cm(0.74)
        ep = bi.add_paragraph()
        ep.paragraph_format.space_after = Pt(2)
        set_font(ep.add_run(r["en"]), "Calibri", 9.5, color=(79, 98, 118))
        cp = bi.add_paragraph(zh)
        cp.paragraph_format.first_line_indent = Cm(0.74)
    out_cn.parent.mkdir(parents=True, exist_ok=True)
    cn.save(out_cn)
    bi.save(out_bi)


def clean_filename(value: str, fallback: str) -> str:
    value = re.sub(r'[<>:"/\\|?*\x00-\x1f]', " ", value)
    value = re.sub(r"\s+", " ", value).strip(" .")
    return (value or fallback)[:90].rstrip(" .")


def infer_title(records: list[dict], pdf: Path) -> str:
    candidates = [r["en"] for r in records if r["page"] == 1 and r["kind"] == "body"]
    first = next((x for x in candidates if 35 <= len(x) <= 300 and not re.search(r"Edited by|Author affiliations|Copyright|Significance", x, re.I)), pdf.stem)
    if "SCIENCES " in first:
        first = first.split("SCIENCES ", 1)[1]
    first = re.sub(r"^.*?RESEARCH ARTICLE\s*\|\s*", "", first, flags=re.I)
    # PNAS title ends before the author line, which is normally a separate block.
    return first.strip() or pdf.stem.replace("-", " ")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("pdf", type=Path)
    ap.add_argument("--output-dir", type=Path)
    ap.add_argument("--argos-runtime", type=Path, default=Path(r"C:\at"))
    ap.add_argument("--argos-models", type=Path, default=Path(r"C:\at\models"))
    ap.add_argument("--final", action="store_true", help="assemble only from fully ai_verified cache rows")
    ap.add_argument("--title-zh", help="human-verified Chinese title for final filenames")
    args = ap.parse_args()
    out = args.output_dir or args.pdf.parent
    work = out / (args.pdf.stem + "_translation_work")
    records_path = work / "source_records.json"
    records = extract_records(args.pdf)
    work.mkdir(parents=True, exist_ok=True)
    records_path.write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
    translator = setup_argos(args.argos_runtime, args.argos_models)
    cache = translate_records(records, work / "translations.jsonl", translator)
    title_en = infer_title(records, args.pdf)
    if args.final:
        unverified = [r["id"] for r in records if r["kind"] == "body" and cache[r["id"]].get("status") != "ai_verified"]
        if unverified:
            raise SystemExit("FINAL ASSEMBLY BLOCKED: non-ai_verified records: " + ", ".join(unverified[:20]))
        if not args.title_zh:
            raise SystemExit("FINAL ASSEMBLY BLOCKED: --title-zh is required; never use a machine-translated title for delivery")
        title_zh = clean_filename(args.title_zh, "论文中文译本")
        # Verified-text staging only. Reserve final Image2 filenames for files
        # that have actually passed figure, navigation, and render QA.
        out_cn = work / f"{title_zh}_AI校订文本_待Image2与链接_全文.docx"
        out_bi = work / f"{title_zh}_AI校订文本_待Image2与链接_逐段对照.docx"
    else:
        title_zh = clean_filename(translator.translate(title_en, "en", "zh").strip(), "论文中文译本")
        out_cn = work / f"{title_zh}_机器初译草稿_全文.docx"
        out_bi = work / f"{title_zh}_机器初译草稿_逐段对照.docx"
    build_docs(records, cache, out_cn, out_bi, title_zh)
    print(json.dumps({
        "status": "final_text_assembly_requires_figure_and_link_postprocessing" if args.final else "machine_draft_not_for_delivery",
        "records": len(records), "cn": str(out_cn), "bilingual": str(out_bi),
        "translation_cache": str(work / "translations.jsonl"),
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()

