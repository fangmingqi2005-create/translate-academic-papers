#!/usr/bin/env python3
"""Add WPS-compatible citation fields, reference bookmarks, and DOI links."""

from __future__ import annotations

import argparse
import copy
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


REFERENCE_HEADING_RE = re.compile(r"^\s*(references|bibliography|参考文献)\s*$", re.I)
REFERENCE_ENTRY_RE = re.compile(r"^\s*(?:\[(\d+)\]|(\d+)[.)])\s+")
DOI_RE = re.compile(r"(?i)(?:https?://(?:dx\.)?doi\.org/)?(10\.\d{4,9}/[-._;()/:A-Z0-9]+)")
CITATION_RE = re.compile(
    r"(?P<open>[（(\[])(?P<content>\s*\d+(?:\s*[-–—,;，；]\s*\d+)*\s*)(?P<close>[）)\]])"
)


def _set_link_style(run_el) -> None:
    rpr = run_el.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run_el.insert(0, rpr)
    color = rpr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rpr.append(color)
    color.set(qn("w:val"), "0563C1")
    underline = rpr.find(qn("w:u"))
    if underline is None:
        underline = OxmlElement("w:u")
        rpr.append(underline)
    underline.set(qn("w:val"), "single")


def _new_run(text: str, template_run=None):
    run_el = OxmlElement("w:r")
    if template_run is not None:
        rpr = template_run.find(qn("w:rPr"))
        if rpr is not None:
            run_el.append(copy.deepcopy(rpr))
    text_el = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        text_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_el.text = text
    run_el.append(text_el)
    return run_el


def _internal_hyperlink(text: str, anchor: str, template_run=None):
    """Return a native OOXML internal hyperlink understood by WPS Writer.

    WPS reliably activates ``w:hyperlink w:anchor`` targets (usually with
    Ctrl+click, depending on the user's WPS hyperlink preference).  A field
    result alone can look blue but is not consistently clickable in WPS.
    """
    field = OxmlElement("w:hyperlink")
    field.set(qn("w:anchor"), anchor)
    field.set(qn("w:history"), "1")
    run_el = _new_run(text, template_run)
    _set_link_style(run_el)
    field.append(run_el)
    return field


def _external_hyperlink(paragraph, text: str, url: str, template_run=None):
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    hyperlink.set(qn("w:history"), "1")
    run_el = _new_run(text, template_run)
    _set_link_style(run_el)
    hyperlink.append(run_el)
    return hyperlink


def _bookmark_paragraph(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _replace_run_with_parts(run, parts) -> None:
    parent = run._r.getparent()
    pos = parent.index(run._r)
    template = run._r
    parent.remove(run._r)
    for element in parts:
        if isinstance(element, str):
            new_element = _new_run(element, template)
        else:
            new_element = element(template)
        parent.insert(pos, new_element)
        pos += 1


def _find_reference_entries(document):
    in_references = False
    entries = {}
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if REFERENCE_HEADING_RE.match(text):
            in_references = True
            continue
        if not in_references:
            continue
        match = REFERENCE_ENTRY_RE.match(text)
        if match:
            number = int(match.group(1) or match.group(2))
            if number in entries:
                print(f"WARNING: duplicate reference number {number}", file=sys.stderr)
            else:
                entries[number] = (paragraph_index, paragraph)
    if entries:
        return entries

    # Some PDF/DOCX builders merge or lose the References heading. Fall back
    # conservatively to a numbered, ascending tail sequence.
    candidates = []
    halfway = len(document.paragraphs) // 2
    for paragraph_index, paragraph in enumerate(document.paragraphs[halfway:], start=halfway):
        match = REFERENCE_ENTRY_RE.match(paragraph.text.strip())
        if match:
            candidates.append((int(match.group(1) or match.group(2)), paragraph_index, paragraph))
    sequence = []
    for number, paragraph_index, paragraph in candidates:
        if not sequence:
            if number == 1:
                sequence = [(number, paragraph_index, paragraph)]
        elif number == sequence[-1][0] + 1:
            sequence.append((number, paragraph_index, paragraph))
        elif number == 1:
            sequence = [(number, paragraph_index, paragraph)]
    if len(sequence) >= 3:
        return {number: (paragraph_index, paragraph) for number, paragraph_index, paragraph in sequence}
    return {}


def _link_body_citations(document, reference_numbers, reference_start_index):
    linked = 0
    unresolved = set()
    max_reference = max(reference_numbers) if reference_numbers else 0

    def split_parts(text):
        nonlocal linked
        parts = []
        cursor = 0
        for match in CITATION_RE.finditer(text):
            parts.append(text[cursor:match.start()])
            parts.append(match.group("open"))
            content = match.group("content")
            number_matches = list(re.finditer(r"\d+", content))
            inner_cursor = 0
            for number_match in number_matches:
                parts.append(content[inner_cursor:number_match.start()])
                number = int(number_match.group())
                if number in reference_numbers:
                    parts.append(lambda template, n=number: _internal_hyperlink(str(n), f"ref_{n}", template))
                    linked += 1
                else:
                    parts.append(number_match.group())
                    if number <= max_reference:
                        unresolved.add(number)
                inner_cursor = number_match.end()
            parts.append(content[inner_cursor:])
            parts.append(match.group("close"))
            cursor = match.end()
        parts.append(text[cursor:])
        return parts

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        if paragraph_index >= reference_start_index:
            break
        paragraph_linked_before = linked
        for run in list(paragraph.runs):
            text = run.text
            if not CITATION_RE.search(text):
                continue
            _replace_run_with_parts(run, split_parts(text))

        # WPS/Word often splits a visible citation such as （1–8）across
        # several runs. Rebuild text-only paragraphs once when no individual
        # run contained a complete citation. Never touch paragraphs with a
        # drawing, because removing their runs would remove the figure.
        if (
            linked == paragraph_linked_before
            and CITATION_RE.search(paragraph.text)
            and paragraph._p.find(".//" + qn("w:drawing")) is None
        ):
            text = paragraph.text
            runs = list(paragraph.runs)
            template = runs[0]._r if runs else None
            for run in runs:
                paragraph._p.remove(run._r)
            for part in split_parts(text):
                element = _new_run(part, template) if isinstance(part, str) else part(template)
                paragraph._p.append(element)
    return linked, unresolved


def _link_dois(reference_paragraphs):
    linked = 0
    for paragraph in reference_paragraphs:
        for run in list(paragraph.runs):
            text = run.text
            matches = list(DOI_RE.finditer(text))
            if not matches:
                continue
            parts = []
            cursor = 0
            for match in matches:
                raw = match.group(0)
                doi = match.group(1).rstrip(".,;:)")
                visible = raw[: len(raw) - (len(match.group(1)) - len(doi))] if len(doi) != len(match.group(1)) else raw
                parts.append(text[cursor:match.start()])
                parts.append(lambda template, v=visible, d=doi: _external_hyperlink(paragraph, v, f"https://doi.org/{d}", template))
                cursor = match.start() + len(visible)
                linked += 1
            parts.append(text[cursor:])
            _replace_run_with_parts(run, parts)
    return linked


def process(input_path: Path, output_path: Path) -> None:
    document = Document(str(input_path))
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    entries = _find_reference_entries(document)
    if not entries:
        raise ValueError("No numbered references found after a References/参考文献 heading")

    for bookmark_id, number in enumerate(sorted(entries), start=1000):
        _bookmark_paragraph(entries[number][1], f"ref_{number}", bookmark_id)

    reference_start_index = min(index for index, _ in entries.values())
    linked_citations, unresolved = _link_body_citations(
        document, set(entries), reference_start_index
    )
    linked_dois = _link_dois(paragraph for _, paragraph in entries.values())

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    print(f"References bookmarked: {len(entries)}")
    print(f"Citation numbers linked: {linked_citations}")
    print(f"DOIs linked: {linked_dois}")
    if unresolved:
        print(
            "WARNING: citation numbers without reference targets: "
            + ", ".join(map(str, sorted(unresolved))),
            file=sys.stderr,
        )


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    if args.input.resolve() == args.output.resolve():
        parser.error("input and output paths must differ")
    process(args.input, args.output)


if __name__ == "__main__":
    main()

